"""Office / 開放文件格式文字擷取。

支援：
- Excel (.xlsx, .xlsm)：逐工作表擷取儲存格
- OpenDocument 試算表 (.ods)：逐工作表
- OpenDocument 文字 (.odt)、Word (.docx)：段落與表格
- PDF (.pdf)：逐頁擷取文字；無文字層時自動走 Azure DI OCR（若已設定）
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Callable, Iterable, List, Optional

from ..settings import MAX_DOCUMENT_ROWS, MAX_DOCUMENT_SHEETS, OCR_FALLBACK_ENABLED
from .spreadsheet_text import rows_to_scan_text

log = logging.getLogger(__name__)

# 二進位文件副檔名 → 擷取函式
XLSX_SUFFIXES = {".xlsx", ".xlsm"}
XLS_SUFFIXES = {".xls"}
EXCEL_SUFFIXES = XLSX_SUFFIXES | XLS_SUFFIXES
ODS_SUFFIXES = {".ods"}
ODT_SUFFIXES = {".odt"}
DOCX_SUFFIXES = {".docx"}
PDF_SUFFIXES = {".pdf"}

DOCUMENT_SUFFIXES = EXCEL_SUFFIXES | ODS_SUFFIXES | ODT_SUFFIXES | DOCX_SUFFIXES | PDF_SUFFIXES


@dataclass(frozen=True)
class DocumentSegment:
    """文件中的一段可掃描文字（例如 Excel 的單一工作表）。"""

    source: str
    text: str


class DocumentReadError(Exception):
    """無法解析文件。"""


def is_document_file(path: str | Path) -> bool:
    return Path(path).suffix.lower() in DOCUMENT_SUFFIXES


def _cell_str(value: object) -> str:
    if value is None or value == "":
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value).strip()
    return str(value).strip()


def _rows_to_text(rows: Iterable[Iterable[object]]) -> str:
    return rows_to_scan_text(rows, max_rows=MAX_DOCUMENT_ROWS)


def _extract_xlsx(data: bytes, filename: str) -> List[DocumentSegment]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentReadError("缺少 openpyxl，無法讀取 Excel") from exc

    try:
        wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:
        raise DocumentReadError(f"無法解析 Excel：{exc}") from exc

    segments: List[DocumentSegment] = []
    try:
        for idx, ws in enumerate(wb.worksheets):
            if idx >= MAX_DOCUMENT_SHEETS:
                break
            text = _rows_to_text(ws.iter_rows(values_only=True))
            if text.strip():
                sheet = ws.title or f"Sheet{idx + 1}"
                segments.append(DocumentSegment(source=f"{filename}#{sheet}", text=text))
    finally:
        wb.close()
    return segments


def _extract_xls(data: bytes, filename: str) -> List[DocumentSegment]:
    try:
        import xlrd
    except ImportError as exc:
        raise DocumentReadError("缺少 xlrd，無法讀取舊版 Excel (.xls)") from exc

    try:
        book = xlrd.open_workbook(file_contents=data)
    except Exception as exc:
        raise DocumentReadError(f"無法解析 Excel (.xls)：{exc}") from exc

    segments: List[DocumentSegment] = []
    for idx in range(book.nsheets):
        if idx >= MAX_DOCUMENT_SHEETS:
            break
        sheet = book.sheet_by_index(idx)
        name = sheet.name or f"Sheet{idx + 1}"

        def _rows():
            for row_idx in range(sheet.nrows):
                if row_idx >= MAX_DOCUMENT_ROWS:
                    yield ["…(已達列數上限)"]
                    return
                yield sheet.row_values(row_idx)

        text = _rows_to_text(_rows())
        if text.strip():
            segments.append(DocumentSegment(source=f"{filename}#{name}", text=text))
    return segments


def _odf_cell_text(cell) -> str:
    from odf.text import P

    parts: List[str] = []
    for p in cell.getElementsByType(P):
        parts.append("".join(node.data for node in p.childNodes if node.nodeType == 3))
    return "".join(parts).strip()


def _extract_ods(data: bytes, filename: str) -> List[DocumentSegment]:
    try:
        from odf.opendocument import load
        from odf.table import Table, TableCell, TableRow
    except ImportError as exc:
        raise DocumentReadError("缺少 odfpy，無法讀取 ODS") from exc

    try:
        doc = load(BytesIO(data))
    except Exception as exc:
        raise DocumentReadError(f"無法解析 ODS：{exc}") from exc

    segments: List[DocumentSegment] = []
    for idx, table in enumerate(doc.getElementsByType(Table)):
        if idx >= MAX_DOCUMENT_SHEETS:
            break
        name = table.getAttribute("name") or f"Sheet{idx + 1}"

        def _rows():
            row_count = 0
            for row in table.getElementsByType(TableRow):
                if row_count >= MAX_DOCUMENT_ROWS:
                    yield ["…(已達列數上限)"]
                    return
                cells = [_odf_cell_text(c) for c in row.getElementsByType(TableCell)]
                yield cells
                row_count += 1

        text = _rows_to_text(_rows())
        if text.strip():
            segments.append(DocumentSegment(source=f"{filename}#{name}", text=text))
    return segments


def _extract_odt(data: bytes, filename: str) -> List[DocumentSegment]:
    try:
        from odf.opendocument import load
        from odf.table import Table, TableCell, TableRow
        from odf.text import P
    except ImportError as exc:
        raise DocumentReadError("缺少 odfpy，無法讀取 ODT") from exc

    try:
        doc = load(BytesIO(data))
    except Exception as exc:
        raise DocumentReadError(f"無法解析 ODT：{exc}") from exc

    parts: List[str] = []
    for p in doc.getElementsByType(P):
        t = "".join(node.data for node in p.childNodes if node.nodeType == 3).strip()
        if t:
            parts.append(t)

    for table in doc.getElementsByType(Table):
        for row in table.getElementsByType(TableRow):
            cells = [_odf_cell_text(c) for c in row.getElementsByType(TableCell)]
            line = "\t".join(c for c in cells if c)
            if line:
                parts.append(line)

    text = "\n".join(parts)
    if not text.strip():
        return []
    return [DocumentSegment(source=filename, text=text)]


def _extract_docx(data: bytes, filename: str) -> List[DocumentSegment]:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentReadError("缺少 python-docx，無法讀取 DOCX") from exc

    try:
        doc = Document(BytesIO(data))
    except Exception as exc:
        raise DocumentReadError(f"無法解析 DOCX：{exc}") from exc

    parts: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    if not text.strip():
        return []
    return [DocumentSegment(source=filename, text=text)]


def _extract_pdf(
    data: bytes,
    filename: str,
    *,
    stats: Optional[dict] = None,
) -> List[DocumentSegment]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentReadError("缺少 pypdf，無法讀取 PDF") from exc

    try:
        reader = PdfReader(BytesIO(data), strict=False)
    except Exception as exc:
        raise DocumentReadError(f"無法解析 PDF：{exc}") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:
            raise DocumentReadError("PDF 已加密，無法讀取") from exc

    pages_text: List[tuple[int, str]] = []
    total_pages = len(reader.pages)
    for idx, page in enumerate(reader.pages):
        if idx >= MAX_DOCUMENT_SHEETS:
            break
        try:
            raw = page.extract_text() or ""
        except Exception:
            raw = ""
        pages_text.append((idx + 1, raw.strip()))

    has_text = any(text for _, text in pages_text)

    if not has_text and OCR_FALLBACK_ENABLED:
        try:
            from ..ocr import OcrError, is_configured, ocr_pdf_pages
        except ImportError:  # pragma: no cover
            OcrError = is_configured = ocr_pdf_pages = None  # type: ignore

        if is_configured and is_configured():
            try:
                ocr_pages = ocr_pdf_pages(data)
            except OcrError as exc:  # type: ignore[misc]
                log.warning("PDF 無文字層且 OCR 失敗：%s", exc)
                if stats is not None:
                    stats["ocr_warning"] = str(exc)
            else:
                pages_text = [
                    (page_no, ocr_pages.get(page_no, "") or existing)
                    for page_no, existing in pages_text
                ]
                if stats is not None:
                    stats["ocr_used"] = True
                    stats["ocr_pages"] = sum(1 for _, t in pages_text if t)
                    stats["ocr_provider"] = "azure-di"
        elif stats is not None:
            stats["ocr_warning"] = (
                "PDF 無文字層；可設定 AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT/KEY 啟用 OCR。"
            )

    segments: List[DocumentSegment] = []
    for page_no, text in pages_text:
        text = (text or "").strip()
        if text:
            segments.append(DocumentSegment(source=f"{filename}#page={page_no}", text=text))

    if not segments:
        raise DocumentReadError(
            "PDF 無可擷取文字；若為掃描影像 PDF，請設定 Azure Document Intelligence 啟用 OCR。"
        )
    if stats is not None:
        stats.setdefault("total_pages", total_pages)
    return segments


_EXTRACTORS: dict[str, Callable[[bytes, str], List[DocumentSegment]]] = {}
for suf in XLSX_SUFFIXES:
    _EXTRACTORS[suf] = _extract_xlsx
for suf in XLS_SUFFIXES:
    _EXTRACTORS[suf] = _extract_xls
for suf in ODS_SUFFIXES:
    _EXTRACTORS[suf] = _extract_ods
for suf in ODT_SUFFIXES:
    _EXTRACTORS[suf] = _extract_odt
for suf in DOCX_SUFFIXES:
    _EXTRACTORS[suf] = _extract_docx
for suf in PDF_SUFFIXES:
    _EXTRACTORS[suf] = _extract_pdf


def supported_document_formats() -> List[str]:
    return sorted(DOCUMENT_SUFFIXES)


def extract_document_segments(
    data: bytes,
    filename: str,
    *,
    ext: str | None = None,
    stats: Optional[dict] = None,
) -> List[DocumentSegment]:
    """依副檔名（或指定 ext）擷取文件各分頁/段落文字。

    若傳入 *stats*，會記錄 OCR 使用情況（``ocr_used``、``ocr_pages`` 等）。
    """
    from .format_sniff import sniff_document_suffix

    chosen = (ext or Path(filename).suffix or "").lower()
    sniffed = sniff_document_suffix(data)
    if sniffed:
        chosen = sniffed
    elif chosen not in DOCUMENT_SUFFIXES:
        chosen = ""
        raise DocumentReadError(
            f"不支援的文件格式 {chosen or '(未知)'}；支援：{', '.join(supported_document_formats())}"
        )
    fn = _EXTRACTORS.get(chosen)
    if fn is None:
        raise DocumentReadError(f"不支援的文件格式 {chosen}")
    try:
        segments = fn(data, filename, stats=stats)  # type: ignore[call-arg]
    except TypeError:
        segments = fn(data, filename)
    if not segments:
        raise DocumentReadError("文件內容為空或無可讀文字")
    return segments
